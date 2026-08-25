import os
import json
import re
import io
import requests
import bleach
from datetime import datetime
from functools import wraps

from flask import (
    Flask, g, jsonify, redirect, render_template, request, session, url_for,
)
from werkzeug.exceptions import HTTPException
from werkzeug.security import check_password_hash, generate_password_hash

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TURSO_URL = os.environ.get("TURSO_DATABASE_URL")
TURSO_TOKEN = os.environ.get("TURSO_AUTH_TOKEN", "")
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
DEFAULT_AI_PROVIDER = os.environ.get("DEFAULT_AI_PROVIDER", "groq")

if TURSO_URL:
    DATABASE_PATH = None
else:
    _db_dir = "/tmp" if not os.access(BASE_DIR, os.W_OK) else BASE_DIR
    DATABASE_PATH = os.path.join(_db_dir, "mazula.db")

app = Flask(__name__)
app.config.update(
    SECRET_KEY=os.environ.get("MAZULA_SECRET_KEY", "mazula-dev-secret-key"),
)

SCHEMA = [
    """CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        email TEXT NOT NULL UNIQUE,
        password TEXT NOT NULL,
        created_at TEXT NOT NULL DEFAULT (datetime('now','localtime'))
    )""",
    """CREATE TABLE IF NOT EXISTS works (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        title TEXT NOT NULL,
        work_type TEXT NOT NULL DEFAULT 'monografia',
        theme TEXT,
        area TEXT,
        keywords TEXT,
        objectives TEXT,
        abstract TEXT DEFAULT '',
        status TEXT NOT NULL DEFAULT 'rascunho',
        target_words INTEGER DEFAULT 10000,
        created_at TEXT NOT NULL DEFAULT (datetime('now','localtime')),
        updated_at TEXT NOT NULL DEFAULT (datetime('now','localtime'))
    )""",
    """CREATE TABLE IF NOT EXISTS sections (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        work_id INTEGER NOT NULL REFERENCES works(id) ON DELETE CASCADE,
        title TEXT NOT NULL,
        content TEXT DEFAULT '',
        order_index INTEGER NOT NULL DEFAULT 0,
        created_at TEXT NOT NULL DEFAULT (datetime('now','localtime')),
        updated_at TEXT NOT NULL DEFAULT (datetime('now','localtime'))
    )""",
    """CREATE TABLE IF NOT EXISTS references_ (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        work_id INTEGER NOT NULL REFERENCES works(id) ON DELETE CASCADE,
        authors TEXT NOT NULL,
        year TEXT,
        title TEXT NOT NULL,
        source TEXT,
        doi TEXT,
        url TEXT,
        ref_type TEXT DEFAULT 'livro',
        pages TEXT,
        publisher TEXT,
        edition TEXT,
        created_at TEXT NOT NULL DEFAULT (datetime('now','localtime'))
    )""",
    """CREATE TABLE IF NOT EXISTS chat_history (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        work_id INTEGER NOT NULL REFERENCES works(id) ON DELETE CASCADE,
        role TEXT NOT NULL,
        content TEXT NOT NULL,
        created_at TEXT NOT NULL DEFAULT (datetime('now','localtime'))
    )""",
    """CREATE TABLE IF NOT EXISTS article_chunks (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        work_id INTEGER NOT NULL REFERENCES works(id) ON DELETE CASCADE,
        ref_id INTEGER REFERENCES references_(id) ON DELETE SET NULL,
        filename TEXT NOT NULL,
        chunk_index INTEGER NOT NULL DEFAULT 0,
        content TEXT NOT NULL,
        created_at TEXT NOT NULL DEFAULT (datetime('now','localtime'))
    )""",
    "CREATE INDEX IF NOT EXISTS idx_works_user ON works(user_id)",
    "CREATE INDEX IF NOT EXISTS idx_sections_work ON sections(work_id, order_index)",
    "CREATE INDEX IF NOT EXISTS idx_references_work ON references_(work_id)",
    "CREATE INDEX IF NOT EXISTS idx_chat_work ON chat_history(work_id)",
    "CREATE INDEX IF NOT EXISTS idx_chunks_work ON article_chunks(work_id)",
]

WORK_TYPES = {
    "monografia": {
        "label": "Monografia",
        "sections": ["Introducao", "Revisao de Literatura", "Fundamentacao Teorica", "Metodologia", "Resultados e Discussao", "Conclusao"],
    },
    "artigo": {
        "label": "Artigo Cientifico",
        "sections": ["Introducao", "Metodos", "Resultados", "Discussao", "Conclusao"],
    },
    "projeto": {
        "label": "Projeto Practico",
        "sections": ["Introducao", "Revisao de Literatura", "Metodologia", "Cronograma", "Referencias"],
    },
    "dissertacao": {
        "label": "Dissertacao",
        "sections": ["Introducao", "Revisao de Literatura", "Fundamentacao Teorica", "Metodologia", "Resultados e Discussao", "Conclusao", "Referencias"],
    },
    "tese": {
        "label": "Tese",
        "sections": ["Introducao", "Revisao de Literatura", "Fundamentacao Teorica", "Metodologia", "Analise e Discussao", "Conclusao", "Referencias"],
    },
    "ensaio": {
        "label": "Ensaio",
        "sections": ["Introducao", "Desenvolvimento", "Conclusao"],
    },
    "revisao": {
        "label": "Revisao de Literatura",
        "sections": ["Introducao", "Metodo de Busca", "Resultados", "Discussao", "Conclusao"],
    },
}


class _TursoRow:
    __slots__ = ("_data", "_cols")
    def __init__(self, data, cols):
        object.__setattr__(self, "_data", data)
        object.__setattr__(self, "_cols", cols)
    def __getitem__(self, key):
        if isinstance(key, int):
            return self._data[key]
        idx = self._cols.index(key)
        return self._data[idx]
    def __contains__(self, key):
        return key in self._cols
    def keys(self):
        return self._cols
    def __iter__(self):
        return iter(self._data)
    def __len__(self):
        return len(self._data)


class _TursoCursor:
    def __init__(self, cursor):
        self._cur = cursor
        self._cols = []
        if hasattr(cursor, "description") and cursor.description:
            self._cols = [d[0] for d in cursor.description]

    @property
    def lastrowid(self):
        return self._cur.lastrowid

    @property
    def rowcount(self):
        return self._cur.rowcount

    def _wrap_row(self, row):
        if row is None or not self._cols:
            return row
        return _TursoRow(row, self._cols)

    def fetchone(self):
        r = self._cur.fetchone()
        return self._wrap_row(r) if r else None

    def fetchall(self):
        return [self._wrap_row(r) for r in self._cur.fetchall()]


class _TursoConn:
    def __init__(self, conn):
        self._conn = conn

    def execute(self, sql, params=None):
        if params is not None:
            params = tuple(params)
            cur = self._conn.execute(sql, params)
        else:
            cur = self._conn.execute(sql)
        cols = []
        if hasattr(cur, "description") and cur.description:
            cols = [d[0] for d in cur.description]
        tc = _TursoCursor(cur)
        tc._cols = cols
        return tc

    def executescript(self, sql):
        for stmt in sql.split(";"):
            stmt = stmt.strip()
            if stmt:
                self.execute(stmt)

    def commit(self):
        self._conn.commit()

    def close(self):
        self._conn.close()


def get_db():
    if "db" not in g:
        if TURSO_URL:
            import libsql_experimental as libsql
            raw = libsql.connect(database=TURSO_URL, auth_token=TURSO_TOKEN)
            g.db = _TursoConn(raw)
        else:
            import sqlite3 as _sqlite3
            conn = _sqlite3.connect(DATABASE_PATH)
            conn.row_factory = _sqlite3.Row
            conn.execute("PRAGMA foreign_keys=ON")
            conn.execute("PRAGMA journal_mode=WAL")
            g.db = conn
    return g.db


@app.teardown_appcontext
def close_db(_exc=None):
    db = g.pop("db", None)
    if db is not None:
        db.close()


def init_db():
    db = get_db()
    for stmt in SCHEMA:
        try:
            db.execute(stmt)
        except Exception:
            pass
    for col, default in [("abstract", ""), ("target_words", 10000)]:
        try:
            db.execute("ALTER TABLE works ADD COLUMN {} DEFAULT {}".format(col, default))
        except Exception:
            pass
    try:
        db.execute("CREATE TABLE IF NOT EXISTS chat_history (id INTEGER PRIMARY KEY AUTOINCREMENT, work_id INTEGER NOT NULL, role TEXT NOT NULL, content TEXT NOT NULL, created_at TEXT)")
    except Exception:
        pass
    try:
        db.execute("UPDATE users SET name = 'Academico', email = 'admin@trabalhofacil.com' WHERE name = 'Nutricionista' OR email = 'admin@nutri.local'")
    except Exception:
        pass
    db.commit()


def row_to_dict(row):
    if row is None:
        return None
    if hasattr(row, "keys"):
        return {k: row[k] for k in row.keys()}
    try:
        return dict(row)
    except Exception:
        return None


def rows_to_list(rows):
    return [row_to_dict(r) for r in rows]


def now_str():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def today_str():
    return datetime.now().strftime("%Y-%m-%d")


def json_body():
    data = request.get_json(silent=True)
    return data if isinstance(data, dict) else {}


ALLOWED_TAGS = ['p', 'br', 'strong', 'em', 'u', 'h3', 'h4', 'blockquote', 'ul', 'ol', 'li', 'a', 'span']
ALLOWED_ATTRS = {'a': ['href', 'target'], 'span': ['style']}


def sanitize_html(text):
    if not text:
        return ""
    clean = bleach.clean(text, tags=ALLOWED_TAGS, attributes=ALLOWED_ATTRS, strip=True)
    return clean


def fetch_owned(table, row_id, user_id):
    db = get_db()
    return db.execute(
        'SELECT * FROM "{}" WHERE id = ? AND user_id = ?'.format(table),
        (row_id, user_id),
    ).fetchone()


def login_required(view):
    @wraps(view)
    def wrapped_view(*args, **kwargs):
        uid = session.get("user_id")
        if uid is None:
            return jsonify({"error": "Sessao expirada. Faca login novamente."}), 401
        db = get_db()
        user = db.execute("SELECT id FROM users WHERE id = ?", (uid,)).fetchone()
        if user is None:
            session.clear()
            return jsonify({"error": "Usuario nao encontrado."}), 401
        return view(*args, **kwargs)
    return wrapped_view


def current_user():
    db = get_db()
    row = db.execute("SELECT * FROM users WHERE id = ?", (session["user_id"],)).fetchone()
    return row_to_dict(row)


def public_user(user):
    if not user:
        return None
    return {"id": user["id"], "name": user["name"], "email": user["email"], "created_at": user.get("created_at")}


def groq_chat(messages, temperature=0.7, max_tokens=4096):
    resp = requests.post(
        "https://api.groq.com/openai/v1/chat/completions",
        headers={"Authorization": "Bearer {}".format(GROQ_API_KEY), "Content-Type": "application/json"},
        json={"model": "qwen/qwen3.6-27b", "messages": messages, "temperature": temperature, "max_tokens": max_tokens},
        timeout=90,
    )
    if resp.status_code != 200:
        error_msg = "Erro ao comunicar com a IA."
        try:
            error_msg = resp.json().get("error", {}).get("message", error_msg)
        except Exception:
            pass
        return None, error_msg
    return resp.json().get("choices", [{}])[0].get("message", {}).get("content", ""), None


def gemini_chat(messages, temperature=0.7, max_tokens=4096):
    contents = []
    system_instruction = None
    for m in messages:
        if m["role"] == "system":
            system_instruction = m["content"]
        else:
            role = "user" if m["role"] == "user" else "model"
            contents.append({"role": role, "parts": [{"text": m["content"]}]})
    body = {"contents": contents, "generationConfig": {"temperature": temperature, "maxOutputTokens": max_tokens}}
    if system_instruction:
        body["systemInstruction"] = {"parts": [{"text": system_instruction}]}
    resp = requests.post(
        "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={}".format(GEMINI_API_KEY),
        headers={"Content-Type": "application/json"},
        json=body,
        timeout=120,
    )
    if resp.status_code != 200:
        error_msg = "Erro ao comunicar com Gemini."
        try:
            error_msg = resp.json().get("error", {}).get("message", error_msg)
        except Exception:
            pass
        return None, error_msg
    try:
        text = resp.json()["candidates"][0]["content"]["parts"][0]["text"]
        return text, None
    except (KeyError, IndexError):
        return None, "Resposta invalida do Gemini."


def ai_chat(messages, provider=None, temperature=0.7, max_tokens=4096):
    provider = provider or DEFAULT_AI_PROVIDER
    if provider == "gemini" and GEMINI_API_KEY:
        return gemini_chat(messages, temperature, max_tokens)
    if provider == "groq" and GROQ_API_KEY:
        return groq_chat(messages, temperature, max_tokens)
    if GROQ_API_KEY:
        return groq_chat(messages, temperature, max_tokens)
    if GEMINI_API_KEY:
        return gemini_chat(messages, temperature, max_tokens)
    return None, "Nenhum provider de IA configurado."


def get_work_context(db, work_id, uid):
    work = db.execute("SELECT * FROM works WHERE id = ? AND user_id = ?", (work_id, uid)).fetchone()
    if not work:
        return None, None
    work = row_to_dict(work)
    ctx = "TRABALHO ACADEMICO:\n"
    ctx += "Titulo: {}\nTema: {}\nArea: {}\nPalavras-chave: {}\nObjectivos: {}\nResumo: {}\n\n".format(
        work.get("title") or "", work.get("theme") or "", work.get("area") or "",
        work.get("keywords") or "", work.get("objectives") or "", work.get("abstract") or "",
    )
    secs = db.execute("SELECT title, content FROM sections WHERE work_id = ? ORDER BY order_index", (work_id,)).fetchall()
    ctx += "SECOES:\n"
    for s in secs:
        sd = row_to_dict(s)
        content = sd.get("content") or ""
        ctx += "\n--- {} ---\n{}\n".format(sd["title"], content[:500] if content else "(vazio)")
    refs = db.execute("SELECT authors, year, title, source FROM references_ WHERE work_id = ? ORDER BY authors", (work_id,)).fetchall()
    if refs:
        ctx += "\nREFERENCIAS:\n"
        for r in refs:
            rd = row_to_dict(r)
            ctx += "- {} ({}) {}. {}\n".format(rd.get("authors") or "", rd.get("year") or "s.d.", rd.get("title") or "", rd.get("source") or "")
    return work, ctx


@app.errorhandler(HTTPException)
def handle_http_exception(exc):
    if request.path.startswith("/api/"):
        return jsonify({"error": exc.description}), exc.code
    return exc


@app.route("/")
@login_required
def index():
    return render_template("index.html")


@app.route("/login")
def login_page():
    return redirect(url_for("index"))


@app.route("/register")
def register_page():
    return redirect(url_for("index"))


@app.route("/api/auth/me")
def auth_me():
    if session.get("user_id") is None:
        return jsonify({"user": None})
    user = current_user()
    if user is None:
        session.clear()
        return jsonify({"user": None})
    return jsonify({"user": public_user(user)})


@app.route("/api/dashboard")
@login_required
def dashboard():
    uid = session["user_id"]
    db = get_db()
    total = db.execute("SELECT COUNT(*) AS c FROM works WHERE user_id = ?", (uid,)).fetchone()["c"]
    rascunho = db.execute("SELECT COUNT(*) AS c FROM works WHERE user_id = ? AND status = 'rascunho'", (uid,)).fetchone()["c"]
    concluido = db.execute("SELECT COUNT(*) AS c FROM works WHERE user_id = ? AND status = 'concluido'", (uid,)).fetchone()["c"]
    recent = rows_to_list(
        db.execute("SELECT * FROM works WHERE user_id = ? ORDER BY updated_at DESC LIMIT 5", (uid,)).fetchall()
    )
    return jsonify({"total": total, "rascunho": rascunho, "concluido": concluido, "recent": recent})


@app.route("/api/work-types")
@login_required
def work_types():
    return jsonify({"types": WORK_TYPES})


@app.route("/api/works", methods=["GET"])
@login_required
def works_list():
    uid = session["user_id"]
    db = get_db()
    rows = db.execute("SELECT * FROM works WHERE user_id = ? ORDER BY updated_at DESC", (uid,)).fetchall()
    return jsonify({"works": rows_to_list(rows)})


@app.route("/api/works", methods=["POST"])
@login_required
def works_create():
    uid = session["user_id"]
    data = json_body()
    title = (data.get("title") or "").strip()
    if not title:
        return jsonify({"error": "Titulo e obrigatorio."}), 400
    work_type_raw = (data.get("work_type") or "monografia").strip()
    TYPE_MAP = {
        "monografia": "monografia", "monograph": "monografia",
        "artigo": "artigo", "artigo científico": "artigo", "artigo cientifico": "artigo", "artigo cientifico": "artigo",
        "projeto": "projeto", "projeto prático": "projeto", "projeto practico": "projeto", "projeto pratico": "projeto",
        "dissertação": "dissertacao", "dissertacao": "dissertacao",
        "tese": "tese",
        "ensaio": "ensaio",
        "revisão de literatura": "revisao", "revisao de literatura": "revisao",
    }
    work_type = TYPE_MAP.get(work_type_raw.lower(), "monografia")
    if work_type not in WORK_TYPES:
        work_type = "monografia"
    db = get_db()
    now = now_str()
    cur = db.execute(
        "INSERT INTO works (user_id, title, work_type, theme, area, keywords, objectives, status, target_words, created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
        (uid, title, work_type, (data.get("theme") or "").strip(), (data.get("area") or "").strip(), (data.get("keywords") or "").strip(), (data.get("objectives") or "").strip(), "rascunho", int(data.get("target_words") or 10000), now, now),
    )
    work_id = cur.lastrowid
    sections = WORK_TYPES.get(work_type, WORK_TYPES["monografia"])["sections"]
    for i, sec_title in enumerate(sections):
        db.execute(
            "INSERT INTO sections (work_id, title, content, order_index, created_at) VALUES (?, ?, ?, ?, ?)",
            (work_id, sec_title, "", i, now),
        )
    db.commit()
    row = db.execute("SELECT * FROM works WHERE id = ? AND user_id = ?", (work_id, uid)).fetchone()
    return jsonify({"work": row_to_dict(row)}), 201


@app.route("/api/works/<int:work_id>", methods=["GET"])
@login_required
def works_get(work_id):
    uid = session["user_id"]
    db = get_db()
    row = db.execute("SELECT * FROM works WHERE id = ? AND user_id = ?", (work_id, uid)).fetchone()
    if row is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    work = row_to_dict(row)
    work["sections"] = rows_to_list(
        db.execute("SELECT * FROM sections WHERE work_id = ? ORDER BY order_index", (work_id,)).fetchall()
    )
    work["references"] = rows_to_list(
        db.execute("SELECT * FROM references_ WHERE work_id = ? ORDER BY authors", (work_id,)).fetchall()
    )
    work["chat_history"] = rows_to_list(
        db.execute("SELECT * FROM chat_history WHERE work_id = ? ORDER BY id", (work_id,)).fetchall()
    )
    total_words = 0
    for s in work["sections"]:
        content = s.get("content") or ""
        s["word_count"] = len(content.split()) if content.strip() else 0
        total_words += s["word_count"]
    work["total_words"] = total_words
    work["article_count"] = db.execute("SELECT COUNT(*) as c FROM article_chunks WHERE work_id = ?", (work_id,)).fetchone()["c"]
    return jsonify({"work": work})


@app.route("/api/works/<int:work_id>", methods=["PUT"])
@login_required
def works_update(work_id):
    uid = session["user_id"]
    db = get_db()
    row = fetch_owned("works", work_id, uid)
    if row is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    data = json_body()
    allowed = ("title", "theme", "area", "keywords", "objectives", "status", "target_words", "abstract")
    updates = {k: data[k] for k in allowed if k in data}
    if not updates:
        return jsonify({"error": "Nenhum campo para atualizar."}), 400
    updates["updated_at"] = now_str()
    set_sql = ", ".join('"{}" = ?'.format(c) for c in updates)
    db.execute(
        'UPDATE works SET {} WHERE id = ? AND user_id = ?'.format(set_sql),
        list(updates.values()) + [work_id, uid],
    )
    db.commit()
    row = db.execute("SELECT * FROM works WHERE id = ? AND user_id = ?", (work_id, uid)).fetchone()
    return jsonify({"work": row_to_dict(row)})


@app.route("/api/works/<int:work_id>", methods=["DELETE"])
@login_required
def works_delete(work_id):
    uid = session["user_id"]
    db = get_db()
    row = fetch_owned("works", work_id, uid)
    if row is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    db.execute("DELETE FROM sections WHERE work_id = ?", (work_id,))
    db.execute("DELETE FROM references_ WHERE work_id = ?", (work_id,))
    db.execute("DELETE FROM chat_history WHERE work_id = ?", (work_id,))
    db.execute("DELETE FROM article_chunks WHERE work_id = ?", (work_id,))
    db.execute("DELETE FROM works WHERE id = ? AND user_id = ?", (work_id, uid))
    db.commit()
    return jsonify({"message": "Trabalho excluido."})


@app.route("/api/works/<int:work_id>/sections", methods=["POST"])
@login_required
def section_create(work_id):
    uid = session["user_id"]
    db = get_db()
    work = fetch_owned("works", work_id, uid)
    if work is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    data = json_body()
    title = (data.get("title") or "Nova Secao").strip()
    content = sanitize_html(data.get("content") or "")
    max_idx = db.execute("SELECT COALESCE(MAX(order_index), -1) as mx FROM sections WHERE work_id = ?", (work_id,)).fetchone()["mx"]
    now = now_str()
    cur = db.execute(
        "INSERT INTO sections (work_id, title, content, order_index, created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?)",
        (work_id, title, content, max_idx + 1, now, now),
    )
    db.execute("UPDATE works SET updated_at = ? WHERE id = ?", (now, work_id))
    db.commit()
    row = db.execute("SELECT * FROM sections WHERE id = ?", (cur.lastrowid,)).fetchone()
    return jsonify({"section": row_to_dict(row)}), 201


@app.route("/api/works/<int:work_id>/sections/<int:section_id>", methods=["DELETE"])
@login_required
def section_delete(work_id, section_id):
    uid = session["user_id"]
    db = get_db()
    work = fetch_owned("works", work_id, uid)
    if work is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    sec = db.execute("SELECT id FROM sections WHERE id = ? AND work_id = ?", (section_id, work_id)).fetchone()
    if sec is None:
        return jsonify({"error": "Secao nao encontrada."}), 404
    db.execute("DELETE FROM sections WHERE id = ?", (section_id,))
    db.execute("UPDATE works SET updated_at = ? WHERE id = ?", (now_str(), work_id))
    db.commit()
    return jsonify({"message": "Secao excluida."})


@app.route("/api/works/<int:work_id>/sections/<int:section_id>", methods=["PUT"])
@login_required
def section_update(work_id, section_id):
    uid = session["user_id"]
    db = get_db()
    work = fetch_owned("works", work_id, uid)
    if work is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    data = json_body()
    content = data.get("content")
    title = data.get("title")
    updates = {"updated_at": now_str()}
    if content is not None:
        updates["content"] = sanitize_html(content)
    if title is not None:
        updates["title"] = title
    set_sql = ", ".join('"{}" = ?'.format(c) for c in updates)
    db.execute(
        'UPDATE sections SET {} WHERE id = ? AND work_id = ?'.format(set_sql),
        list(updates.values()) + [section_id, work_id],
    )
    db.execute("UPDATE works SET updated_at = ? WHERE id = ?", (now_str(), work_id))
    db.commit()
    row = db.execute("SELECT * FROM sections WHERE id = ?", (section_id,)).fetchone()
    return jsonify({"section": row_to_dict(row)})


@app.route("/api/works/<int:work_id>/sections/bulk", methods=["PUT"])
@login_required
def sections_bulk_update(work_id):
    uid = session["user_id"]
    db = get_db()
    work = fetch_owned("works", work_id, uid)
    if work is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    data = json_body()
    sections = data.get("sections") or []
    if not sections:
        return jsonify({"error": "Nenhuma secção para atualizar."}), 400
    now = now_str()
    updated = 0
    for s in sections:
        sid = s.get("id")
        if not sid:
            continue
        content = sanitize_html(s.get("content") or "")
        title = s.get("title")
        updates = {"updated_at": now}
        if content is not None:
            updates["content"] = content
        if title is not None:
            updates["title"] = title
        set_sql = ", ".join('"{}" = ?'.format(c) for c in updates)
        db.execute(
            'UPDATE sections SET {} WHERE id = ? AND work_id = ?'.format(set_sql),
            list(updates.values()) + [sid, work_id],
        )
        updated += 1
    if updated:
        db.execute("UPDATE works SET updated_at = ?, word_count = ? WHERE id = ?", (now, data.get("word_count", 0), work_id))
    db.commit()
    return jsonify({"updated": updated})


@app.route("/api/works/<int:work_id>/references", methods=["GET"])
@login_required
def refs_list(work_id):
    uid = session["user_id"]
    db = get_db()
    work = fetch_owned("works", work_id, uid)
    if work is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    rows = db.execute("SELECT * FROM references_ WHERE work_id = ? ORDER BY authors", (work_id,)).fetchall()
    return jsonify({"references": rows_to_list(rows)})


@app.route("/api/works/<int:work_id>/references", methods=["POST"])
@login_required
def refs_create(work_id):
    uid = session["user_id"]
    db = get_db()
    work = fetch_owned("works", work_id, uid)
    if work is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    data = json_body()
    authors = (data.get("authors") or "").strip()
    title = (data.get("title") or "").strip()
    if not authors or not title:
        return jsonify({"error": "Autores e titulo sao obrigatorios."}), 400
    now = now_str()
    cur = db.execute(
        "INSERT INTO references_ (work_id, authors, year, title, source, doi, url, ref_type, pages, publisher, edition, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
        (work_id, authors, (data.get("year") or "").strip(), title, (data.get("source") or "").strip(), (data.get("doi") or "").strip(), (data.get("url") or "").strip(), (data.get("ref_type") or "livro").strip(), (data.get("pages") or "").strip(), (data.get("publisher") or "").strip(), (data.get("edition") or "").strip(), now),
    )
    db.commit()
    row = db.execute("SELECT * FROM references_ WHERE id = ?", (cur.lastrowid,)).fetchone()
    return jsonify({"reference": row_to_dict(row)}), 201


@app.route("/api/works/<int:work_id>/references/<int:ref_id>", methods=["DELETE"])
@login_required
def refs_delete(work_id, ref_id):
    uid = session["user_id"]
    db = get_db()
    work = fetch_owned("works", work_id, uid)
    if work is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    db.execute("DELETE FROM references_ WHERE id = ? AND work_id = ?", (ref_id, work_id))
    db.commit()
    return jsonify({"message": "Referencia excluida."})


@app.route("/api/works/<int:work_id>/references/format-apa", methods=["GET"])
@login_required
def refs_format_apa(work_id):
    uid = session["user_id"]
    db = get_db()
    work = fetch_owned("works", work_id, uid)
    if work is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    rows = db.execute("SELECT * FROM references_ WHERE work_id = ? ORDER BY authors", (work_id,)).fetchall()
    refs = rows_to_list(rows)
    formatted = []
    for r in refs:
        authors = (r.get("authors") or "").strip().rstrip(".")
        year = r.get("year") or "s.d."
        title = (r.get("title") or "").strip().rstrip(".")
        source = (r.get("source") or "").strip()
        doi = (r.get("doi") or "").strip()
        pages = (r.get("pages") or "").strip()
        publisher = (r.get("publisher") or "").strip()
        edition = (r.get("edition") or "").strip()
        apa = "{} ({})".format(authors, year)
        apa += ". {}.".format(title)
        if source:
            if pages:
                apa += " {}. pp. {}.".format(source, pages)
            else:
                apa += " {}.".format(source)
        if edition:
            apa += " {} ed.".format(edition)
        if publisher:
            apa += " {}.".format(publisher)
        if doi:
            apa += " https://doi.org/{}".format(doi)
        elif not source:
            apa += "."
        formatted.append({"id": r["id"], "apa": apa})
    return jsonify({"formatted": formatted})


@app.route("/api/works/<int:work_id>/chat", methods=["GET"])
@login_required
def chat_history(work_id):
    uid = session["user_id"]
    db = get_db()
    work = fetch_owned("works", work_id, uid)
    if work is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    rows = db.execute("SELECT * FROM chat_history WHERE work_id = ? ORDER BY id", (work_id,)).fetchall()
    return jsonify({"messages": rows_to_list(rows)})


@app.route("/api/works/<int:work_id>/chat", methods=["DELETE"])
@login_required
def chat_clear(work_id):
    uid = session["user_id"]
    db = get_db()
    work = fetch_owned("works", work_id, uid)
    if work is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    db.execute("DELETE FROM chat_history WHERE work_id = ?", (work_id,))
    db.commit()
    return jsonify({"message": "Historico limpo."})


@app.route("/api/all")
@login_required
def api_all():
    uid = session["user_id"]
    db = get_db()
    works = rows_to_list(db.execute("SELECT * FROM works WHERE user_id = ? ORDER BY updated_at DESC", (uid,)).fetchall())
    refs = rows_to_list(db.execute("SELECT * FROM references_ WHERE work_id IN (SELECT id FROM works WHERE user_id = ?) ORDER BY authors", (uid,)).fetchall())
    chat = rows_to_list(db.execute("SELECT * FROM chat_history WHERE work_id IN (SELECT id FROM works WHERE user_id = ?) ORDER BY id", (uid,)).fetchall())
    total_wc = 0
    for w in works:
        secs = db.execute("SELECT content FROM sections WHERE work_id = ?", (w["id"],)).fetchall()
        wc = 0
        for s in secs:
            c = s[0] if isinstance(s, tuple) else s["content"]
            wc += len(c.split()) if c and c.strip() else 0
        w["word_count"] = wc
        total_wc += wc
        w["article_count"] = db.execute("SELECT COUNT(*) as c FROM article_chunks WHERE work_id = ?", (w["id"],)).fetchone()["c"]
    return jsonify({"works": works, "references": refs, "chat": chat, "total_words": total_wc})


@app.route("/api/ai/providers", methods=["GET"])
@login_required
def api_ai_providers():
    return jsonify({
        "providers": {
            "groq": {"name": "Groq", "available": bool(GROQ_API_KEY), "free": True, "model": "Qwen 3.6 27B"},
            "gemini": {"name": "Google Gemini", "available": bool(GEMINI_API_KEY), "free": True, "model": "Gemini 2.0 Flash"},
        },
        "default": DEFAULT_AI_PROVIDER,
    })


@app.route("/api/ai", methods=["POST"])
@login_required
def api_ai():
    if not GROQ_API_KEY and not GEMINI_API_KEY:
        return jsonify({"error": "Nenhum provider de IA configurado."}), 503
    data = json_body()
    mode = data.get("mode") or "generate"
    prompt = (data.get("prompt") or "").strip()
    work_id = data.get("work_id")
    lang = data.get("language") or "Portugues"
    tone = data.get("tone") or "Academico"
    provider = data.get("provider") or None
    uid = session["user_id"]
    db = get_db()
    work_ctx = ""
    if work_id:
        _, work_ctx = get_work_context(db, work_id, uid) or ("", "")
    sys_base = """Voce e um assistente academico especializado em trabalhos cientificos em lingua portuguesa de Mocambique.
REGRAS: linguagem formal e academica, terceira pessoa, vocabulario tecnico, sem abreviacoes informais, texto original.
Idioma: {}. Tom: {}.""".format(lang, tone)
    if mode == "generate":
        sys_msg = sys_base + "\n\nGere conteudo academico completo."
        if work_ctx:
            sys_msg += "\n\nContexto:\n" + work_ctx[:1500]
        sys_msg += "\n\nOrientacao: {}".format(prompt or "Gere texto academico completo (minimo 300 palavras).")
    elif mode == "summarize":
        sys_msg = sys_base + "\n\nResuma o seguinte texto academico:\n\n" + (prompt or "")
    elif mode == "expand":
        sys_msg = sys_base + "\n\nExpanda o seguinte texto academico:\n\n" + (prompt or "")
    elif mode == "rewrite":
        sys_msg = sys_base + "\n\nReescreva o seguinte texto academico de forma melhorada:\n\n" + (prompt or "")
    elif mode == "translate":
        sys_msg = "Traduza o texto academico seguinte para {}. Mantenha o formato academico.\n\n{}".format(lang, prompt or "")
    elif mode == "correct":
        sys_msg = sys_base + "\n\nCorrija erros de ortografia, gramatica e estilo:\n\n" + (prompt or "")
    else:
        sys_msg = sys_base
        if work_ctx:
            sys_msg += "\n\nContexto:\n" + work_ctx[:1000]
        sys_msg += "\n\n" + (prompt or "")
    if mode == "chat" and work_id:
        db.execute("INSERT INTO chat_history (work_id, role, content, created_at) VALUES (?, 'user', ?, ?)", (work_id, prompt, now_str()))
        db.commit()
    messages = [{"role": "system", "content": sys_msg}, {"role": "user", "content": prompt or "Ajude-me."}]
    text, error = ai_chat(messages, provider=provider)
    if error:
        return jsonify({"error": error}), 502
    if mode == "chat" and work_id:
        db.execute("INSERT INTO chat_history (work_id, role, content, created_at) VALUES (?, 'assistant', ?, ?)", (work_id, text, now_str()))
        db.commit()
    return jsonify({"result": text, "mode": mode})


@app.route("/api/import_ref", methods=["POST"])
@login_required
def api_import_ref():
    data = json_body()
    identifier = (data.get("identifier") or "").strip()
    if not identifier:
        return jsonify({"error": "Forneca um DOI ou ISBN."}), 400
    ref_data = None
    if identifier.startswith("10."):
        try:
            r = requests.get("https://api.crossref.org/works/" + identifier.lstrip("doi:").lstrip("DOI:").strip(), timeout=15)
            if r.status_code == 200:
                d = r.json().get("message", {})
                authors_list = d.get("author", [])
                if authors_list:
                    parts = []
                    for a in authors_list:
                        family = a.get("family", "")
                        given = a.get("given", "")
                        if family and given:
                            parts.append("{}, {}".format(family, given[0] + "."))
                        elif family:
                            parts.append(family)
                    authors_str = ", ".join(parts)
                else:
                    authors_str = "Autor desconhecido"
                year = ""
                dp = d.get("published-print") or d.get("published-online") or d.get("created")
                if dp and dp.get("date-parts"):
                    year = str(dp["date-parts"][0][0])
                source = ""
                container = d.get("container-title")
                if container:
                    source = container[0] if isinstance(container, list) else container
                ref_data = {
                    "authors": authors_str,
                    "year": year,
                    "title": (d.get("title") or [""])[0] if d.get("title") else "",
                    "source": source,
                    "doi": identifier.lstrip("doi:").lstrip("DOI:").strip(),
                    "url": "https://doi.org/" + identifier.lstrip("doi:").lstrip("DOI:").strip(),
                    "ref_type": "artigo",
                    "publisher": d.get("publisher", ""),
                }
        except Exception:
            pass
    if not ref_data:
        try:
            r = requests.get("https://openlibrary.org/isbn/{}.json".format(identifier), timeout=15)
            if r.status_code == 200:
                d = r.json()
                title = d.get("title", "")
                authors_str = "Autor desconhecido"
                if d.get("authors"):
                    a_ids = [a.get("key") for a in d["authors"] if a.get("key")]
                    a_names = []
                    for aid in a_ids[:5]:
                        try:
                            ar = requests.get("https://openlibrary.org{}.json".format(aid), timeout=10)
                            if ar.status_code == 200:
                                ad = ar.json()
                                name = ad.get("name", "")
                                if name:
                                    a_names.append(name)
                        except Exception:
                            pass
                    if a_names:
                        authors_str = ", ".join(a_names)
                year = ""
                if d.get("publish_date"):
                    m = re.search(r"(\d{4})", d["publish_date"])
                    if m:
                        year = m.group(1)
                publishers = d.get("publishers") or []
                source = publishers[0] if publishers else ""
                ref_data = {
                    "authors": authors_str,
                    "year": year,
                    "title": title,
                    "source": source,
                    "doi": "",
                    "url": "https://openlibrary.org/isbn/{}".format(identifier),
                    "ref_type": "livro",
                    "publisher": source,
                }
        except Exception:
            pass
    if not ref_data:
        return jsonify({"error": "Nao foi possivel encontrar dados para '{}'.".format(identifier)}), 404
    return jsonify(ref_data)


@app.route("/api/chat", methods=["POST"])
@login_required
def api_chat():
    if not GROQ_API_KEY and not GEMINI_API_KEY:
        return jsonify({"error": "Nenhum provider de IA configurado."}), 503
    data = json_body()
    msg = (data.get("message") or "").strip()
    work_id = data.get("work_id")
    provider = data.get("provider") or None
    if not msg:
        return jsonify({"error": "Mensagem obrigatoria."}), 400
    uid = session["user_id"]
    db = get_db()
    work_ctx = data.get("context") or ""
    if work_id and not work_ctx:
        _, work_ctx = get_work_context(db, work_id, uid) or ("", "")
    sys_msg = """Voce e um assistente academico especializado em trabalhos cientificos em Mocambique.
REGRAS: linguagem formal e academica, terceira pessoa, vocabulario tecnico, sem abreviacoes informais, texto original, cite fontes reais APA 7a.
Se o usuario pedir para gerar conteudo, gere paragrafos academicos completos com citações APA.
Seja directo, util e preciso. Evite respostas genericas."""
    if work_ctx:
        sys_msg += "\n\nContexto do trabalho:\n" + work_ctx[:2000]
    if work_id:
        db.execute("INSERT INTO chat_history (work_id, role, content, created_at) VALUES (?, 'user', ?, ?)", (work_id, msg, now_str()))
        db.commit()
        history = db.execute("SELECT role, content FROM chat_history WHERE work_id = ? ORDER BY id DESC LIMIT 20", (work_id,)).fetchall()
        messages = [{"role": "system", "content": sys_msg}]
        for h in reversed(history):
            hd = row_to_dict(h)
            role = hd["role"] if hd["role"] in ("user", "assistant") else "user"
            messages.append({"role": role, "content": hd["content"][:1000]})
    else:
        messages = [{"role": "system", "content": sys_msg}, {"role": "user", "content": msg}]
    text, error = ai_chat(messages, provider=provider)
    if error:
        return jsonify({"error": error}), 502
    if work_id:
        db.execute("INSERT INTO chat_history (work_id, role, content, created_at) VALUES (?, 'assistant', ?, ?)", (work_id, text, now_str()))
        db.commit()
    return jsonify({"result": text})


@app.route("/api/references", methods=["POST"])
@login_required
def api_refs_create():
    uid = session["user_id"]
    data = json_body()
    work_id = data.get("work_id")
    if work_id:
        db = get_db()
        work = fetch_owned("works", work_id, uid)
        if work is None:
            return jsonify({"error": "Trabalho nao encontrado."}), 404
    authors = (data.get("authors") or "").strip()
    title = (data.get("title") or "").strip()
    if not authors or not title:
        return jsonify({"error": "Autores e titulo obrigatorios."}), 400
    db = get_db()
    now = now_str()
    cur = db.execute(
        "INSERT INTO references_ (work_id, authors, year, title, source, doi, url, ref_type, pages, publisher, edition, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
        (work_id or 0, authors, (data.get("year") or "").strip(), title, (data.get("source") or "").strip(), (data.get("doi") or "").strip(), (data.get("url") or "").strip(), (data.get("ref_type") or "livro").strip(), (data.get("pages") or "").strip(), (data.get("publisher") or "").strip(), (data.get("edition") or "").strip(), now),
    )
    db.commit()
    row = db.execute("SELECT * FROM references_ WHERE id = ?", (cur.lastrowid,)).fetchone()
    d = row_to_dict(row)
    authors_apa = (d.get("authors") or "").strip().rstrip(".")
    year_apa = d.get("year") or "s.f."
    title_apa = (d.get("title") or "").strip().rstrip(".")
    d["citation_apa"] = "{} ({}). {}.".format(authors_apa, year_apa, title_apa)
    return jsonify(d), 201


@app.route("/api/references/<int:ref_id>", methods=["DELETE"])
@login_required
def api_refs_delete(ref_id):
    uid = session["user_id"]
    db = get_db()
    ref = db.execute(
        "SELECT r.id FROM references_ r JOIN works w ON r.work_id = w.id WHERE r.id = ? AND w.user_id = ?",
        (ref_id, uid)
    ).fetchone()
    if ref is None:
        return jsonify({"error": "Referencia nao encontrada."}), 404
    db.execute("DELETE FROM references_ WHERE id = ?", (ref_id,))
    db.commit()
    return jsonify({"message": "Referencia removida."})


@app.route("/api/upload-article", methods=["POST"])
@login_required
def upload_article():
    from PyPDF2 import PdfReader
    work_id = request.form.get("work_id")
    if not work_id:
        return jsonify({"error": "work_id obrigatorio."}), 400
    uid = session["user_id"]
    db = get_db()
    work = fetch_owned("works", int(work_id), uid)
    if work is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    if "file" not in request.files:
        return jsonify({"error": "Nenhum ficheiro enviado."}), 400
    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "Nome de ficheiro vazio."}), 400
    ext = f.filename.rsplit(".", 1)[-1].lower() if "." in f.filename else ""
    if ext not in ("pdf", "txt"):
        return jsonify({"error": "Formato nao suportado. Use PDF ou TXT."}), 400
    try:
        if ext == "pdf":
            reader = PdfReader(f)
            text = ""
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        else:
            text = f.read().decode("utf-8", errors="replace")
    except Exception as e:
        return jsonify({"error": "Erro ao ler ficheiro: {}".format(str(e))}), 500
    if not text.strip():
        return jsonify({"error": "Nenhum texto extraido do ficheiro."}), 400
    title_from_pdf = ""
    first_lines = text.strip()[:200]
    if first_lines:
        title_from_pdf = first_lines.split("\n")[0].strip()[:200]
    authors_str = "Autor desconhecido"
    year = ""
    now = now_str()
    cur_ref = db.execute(
        "INSERT INTO references_ (work_id, authors, year, title, source, doi, url, ref_type, pages, publisher, edition, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
        (int(work_id), authors_str, year, title_from_pdf or f.filename, f.filename, "", "", "artigo", "", "", "", now),
    )
    ref_id = cur_ref.lastrowid
    db.commit()
    chunk_size = 800
    overlap = 200
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk = " ".join(words[i:i + chunk_size])
        if chunk.strip():
            chunks.append(chunk)
        i += chunk_size - overlap
    for idx, chunk in enumerate(chunks):
        db.execute(
            "INSERT INTO article_chunks (work_id, ref_id, filename, chunk_index, content, created_at) VALUES (?, ?, ?, ?, ?, ?)",
            (int(work_id), ref_id, f.filename, idx, chunk, now),
        )
    db.commit()
    return jsonify({
        "message": "Artigo importado com sucesso.",
        "ref_id": ref_id,
        "filename": f.filename,
        "chunks": len(chunks),
        "chars": len(text),
    }), 201


@app.route("/api/works/<int:work_id>/articles", methods=["GET"])
@login_required
def list_articles(work_id):
    uid = session["user_id"]
    db = get_db()
    work = fetch_owned("works", work_id, uid)
    if work is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    rows = db.execute(
        "SELECT r.id, r.title, r.authors, r.year, COUNT(c.id) as chunks FROM references_ r LEFT JOIN article_chunks c ON c.ref_id = r.id WHERE r.work_id = ? GROUP BY r.id ORDER BY r.title",
        (work_id,)
    ).fetchall()
    return jsonify({"articles": rows_to_list(rows)})


@app.route("/api/works/<int:work_id>/articles/<int:ref_id>", methods=["DELETE"])
@login_required
def delete_article(work_id, ref_id):
    uid = session["user_id"]
    db = get_db()
    work = fetch_owned("works", work_id, uid)
    if work is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    db.execute("DELETE FROM article_chunks WHERE ref_id = ? AND work_id = ?", (ref_id, work_id))
    db.execute("DELETE FROM references_ WHERE id = ? AND work_id = ?", (ref_id, work_id))
    db.commit()
    return jsonify({"message": "Artigo removido."})


@app.route("/api/works/<int:work_id>/generate-from-articles", methods=["POST"])
@login_required
def generate_from_articles(work_id):
    if not GROQ_API_KEY and not GEMINI_API_KEY:
        return jsonify({"error": "Nenhum provider de IA configurado."}), 503
    uid = session["user_id"]
    db = get_db()
    work = fetch_owned("works", work_id, uid)
    if work is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    data = json_body()
    prompt = (data.get("prompt") or "").strip()
    section_title = (data.get("section_title") or "").strip()
    provider = data.get("provider") or None
    if not prompt:
        prompt = "Escreva o conteudo da secção '{}'".format(section_title or "Trabalho")
    all_chunks = db.execute(
        "SELECT content, filename FROM article_chunks WHERE work_id = ? ORDER BY chunk_index LIMIT 50",
        (work_id,)
    ).fetchall()
    if not all_chunks:
        return jsonify({"error": "Nenhum artigo importado. Carregue artigos PDF primeiro."}), 400
    refs = db.execute(
        "SELECT authors, year, title, doi FROM references_ WHERE work_id = ? ORDER BY authors",
        (work_id,)
    ).fetchall()
    refs_text = ""
    for r in refs:
        rd = row_to_dict(r)
        refs_text += "- {} ({}) {}. {}\n".format(
            rd.get("authors") or "", rd.get("year") or "s.d.",
            rd.get("title") or "", rd.get("doi") or ""
        )
    chunks_text = ""
    for c in all_chunks:
        cd = row_to_dict(c)
        chunks_text += "\n--- Fonte: {} ---\n{}\n".format(cd["filename"], cd["content"][:600])
    sys_msg = """Voce e um assistente academico especializado em trabalhos cientificos em Mocambique.
REGRAS OBRIGATORIAS:
1. Use APENAS informacoes dos artigos fornecidos abaixo como fonte.
2. Cite sempre que possivel usando formato APA 7a: (Autores, Ano)
3. NUNCA invente autores, datas ou dados. Use apenas o que esta nos artigos.
4. Linguagem formal, academica, terceira pessoa, vocabulario tecnico.
5. Paragrafos completos com no minimo 3 frases cada.
6. Inclua citacoes no texto e referencie os artigos citados.

ARTIGOS DISPONIVEIS:
{}
REFERENCIAS:
{}
""".format(chunks_text[:6000], refs_text)
    messages = [
        {"role": "system", "content": sys_msg},
        {"role": "user", "content": prompt + "\n\nEscreva o conteudo academico baseado nos artigos acima. Cite em APA 7a."},
    ]
    text, error = ai_chat(messages, provider=provider, temperature=0.6, max_tokens=4096)
    if error:
        return jsonify({"error": error}), 502
    return jsonify({"result": text})


@app.route("/api/upload-data", methods=["POST"])
@login_required
def upload_data():
    if "file" not in request.files:
        return jsonify({"error": "Nenhum ficheiro enviado."}), 400
    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "Nome de ficheiro vazio."}), 400
    ext = f.filename.rsplit(".", 1)[-1].lower() if "." in f.filename else ""
    if ext not in ("csv", "xlsx", "xls"):
        return jsonify({"error": "Formato nao suportado. Use CSV ou Excel."}), 400
    try:
        if ext == "csv":
            import csv, io
            content = f.read().decode("utf-8", errors="replace")
            reader = csv.reader(io.StringIO(content))
            rows = [row for row in reader if any(cell.strip() for cell in row)]
        else:
            import openpyxl
            wb = openpyxl.load_workbook(f, read_only=True)
            ws = wb.active
            rows = []
            for row in ws.iter_rows(values_only=True):
                r = [str(cell) if cell is not None else "" for cell in row]
                if any(c.strip() for c in r):
                    rows.append(r)
            wb.close()
    except Exception as e:
        return jsonify({"error": "Erro ao ler ficheiro: {}".format(str(e))}), 500
    if len(rows) < 2:
        return jsonify({"error": "Ficheiro precisa de pelo menos 2 linhas (cabecalho + dados)."}), 400
    headers = rows[0]
    data_rows = rows[1:]
    return jsonify({
        "filename": f.filename,
        "headers": headers,
        "rows": data_rows[:100],
        "total_rows": len(data_rows),
    })


@app.route("/api/analyze-data", methods=["POST"])
@login_required
def analyze_data():
    if not GROQ_API_KEY:
        return jsonify({"error": "Chave API do Groq nao configurada."}), 503
    data = json_body()
    headers = data.get("headers") or []
    rows = data.get("rows") or []
    prompt = (data.get("prompt") or "").strip()
    norm = data.get("norm") or "APA"
    provider = data.get("provider") or None
    if not headers or not rows:
        return jsonify({"error": "Dados obrigatorios."}), 400
    table_str = " | ".join(headers) + "\n" + " | ".join(["---"] * len(headers)) + "\n"
    for row in rows[:50]:
        table_str += " | ".join(str(c)[:50] for c in row) + "\n"
    sys_msg = """Voce e um analista de dados academicos especializado em estatistica descritiva.
Analise os dados fornecidos e gere:
1. Tabelas formatadas em {norm} com estatisticas descritivas (media, mediana, moda, DP, min, max)
2. Analise interpretativa dos dados em linguagem academica
3. Formate tabelas usando o padrao academico (sem linhas verticais, apenas horizontal)

Dados:
{dados}
""".format(norm=norm, dados=table_str[:4000])
    user_msg = prompt or "Analise estes dados e gere tabelas em formato academico {} com estatisticas descritivas e interpretação.".format(norm)
    messages = [
        {"role": "system", "content": sys_msg},
        {"role": "user", "content": user_msg},
    ]
    text, error = ai_chat(messages, provider=provider, temperature=0.5, max_tokens=4096)
    if error:
        return jsonify({"error": error}), 502
    return jsonify({"result": text})


@app.route("/api/generate-questionnaire", methods=["POST"])
@login_required
def generate_questionnaire():
    if not GROQ_API_KEY and not GEMINI_API_KEY:
        return jsonify({"error": "Nenhum provider de IA configurado."}), 503
    uid = session["user_id"]
    data = json_body()
    work_id = data.get("work_id")
    prompt = (data.get("prompt") or "").strip()
    provider = data.get("provider") or None
    db = get_db()
    work_ctx = ""
    if work_id:
        _, work_ctx = get_work_context(db, int(work_id), uid) or ("", "")
    sys_msg = """Voce e um metodologo de pesquisa academica especializado em elaborar instrumentos de coleta de dados.
Gere um questionario completo e profissional em formato academico.
REGRAS:
1. Cabecalho: titulo, instituicao, data, instrucoes gerais
2. Secoes claras com titulo
3. Questoes variadas: escala Likert, escolha multipla, aberta, sim/nao
4. Linguagem clara, neutra, sem ambiguidades
5. Minimo 15 questoes distribuidas por seccoes
6. Formato numerado por secao
7. Incluir escala de Likert completa quando aplicavel

{}
""".format(work_ctx[:2000] if work_ctx else "Contexto: Trabalho academico generico.")
    user_msg = prompt or "Gere um instrumento de coleta de dados (questionario) completo para este trabalho."
    messages = [
        {"role": "system", "content": sys_msg},
        {"role": "user", "content": user_msg},
    ]
    text, error = ai_chat(messages, provider=provider, temperature=0.6, max_tokens=4096)
    if error:
        return jsonify({"error": error}), 502
    return jsonify({"result": text})


@app.route("/api/login", methods=["POST"])
def api_login():
    data = json_body()
    email = (data.get("email") or "").strip().lower()
    password = data.get("password") or ""
    if not email or not password:
        return jsonify({"error": "Email e senha obrigatorios."}), 400
    db = get_db()
    user = db.execute("SELECT * FROM users WHERE email = ?", (email,)).fetchone()
    if user is None or not check_password_hash(user["password"], password):
        return jsonify({"error": "Email ou senha invalidos."}), 401
    session.clear()
    session["user_id"] = user["id"]
    return jsonify({"name": user["name"], "email": user["email"], "id": user["id"]})


@app.route("/api/register", methods=["POST"])
def api_register():
    data = json_body()
    name = (data.get("name") or "").strip()
    email = (data.get("email") or "").strip().lower()
    password = data.get("password") or ""
    if not name or not email or not password:
        return jsonify({"error": "Nome, email e senha obrigatorios."}), 400
    if len(password) < 6:
        return jsonify({"error": "A senha deve ter pelo menos 6 caracteres."}), 400
    db = get_db()
    existing = db.execute("SELECT id FROM users WHERE email = ?", (email,)).fetchone()
    if existing is not None:
        return jsonify({"error": "Ja existe uma conta com este email."}), 409
    cur = db.execute(
        "INSERT INTO users (name, email, password, created_at) VALUES (?, ?, ?, ?)",
        (name, email, generate_password_hash(password), now_str()),
    )
    db.commit()
    session.clear()
    session["user_id"] = cur.lastrowid
    return jsonify({"name": name, "email": email, "id": cur.lastrowid}), 201


@app.route("/api/export/<int:work_id>")
@login_required
def export_work(work_id):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT

    uid = session["user_id"]
    db = get_db()
    work, ctx = get_work_context(db, work_id, uid)
    if work is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    sections = rows_to_list(
        db.execute("SELECT * FROM sections WHERE work_id = ? ORDER BY order_index", (work_id,)).fetchall()
    )
    refs = rows_to_list(
        db.execute("SELECT * FROM references_ WHERE work_id = ? ORDER BY authors", (work_id,)).fetchall()
    )
    total_words = 0
    for s in sections:
        c = s.get("content") or ""
        s["word_count"] = len(c.split()) if c.strip() else 0
        total_words += s["word_count"]

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)

    title = work.get("title", "Trabalho Academico")

    p_institution = doc.add_paragraph()
    p_institution.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_institution.add_run("UNIVERSIDADE / INSTITUICAO")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    if work.get("theme"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Tema: ")
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        run2 = p.add_run(work["theme"])
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)

    if work.get("area"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Area: ")
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        run2 = p.add_run(work["area"])
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_author.add_run("Autor: {}".format(session.get("user_name", "Academico")))
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_date.add_run(datetime.now().strftime("%Y"))
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    doc.add_page_break()

    toc_heading = doc.add_heading("SUMARIO", level=1)
    for run in toc_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
    for i, s in enumerate(sections):
        content = s.get("content") or ""
        if content.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run("{}. {}".format(i + 1, s["title"]))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)

    doc.add_page_break()

    if work.get("keywords"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Palavras-chave: ")
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        run2 = p.add_run(work["keywords"])
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        doc.add_page_break()

    for s in sections:
        h = doc.add_heading(s["title"], level=1)
        for run in h.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

        content = s.get("content") or ""
        clean = re.sub(r'<[^>]+>', ' ', content)
        clean = re.sub(r'\s+', ' ', clean).strip()
        paragraphs = [p.strip() for p in clean.split('\n') if p.strip()]

        for para_text in paragraphs:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.line_spacing = 2.0
            run = p.add_run(para_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    if refs:
        doc.add_page_break()
        h = doc.add_heading("Referencias Bibliograficas", level=1)
        for run in h.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)

        for r in refs:
            authors = (r.get("authors") or "").strip().rstrip(".")
            year = r.get("year") or "s.d."
            title_ref = (r.get("title") or "").strip().rstrip(".")
            source = (r.get("source") or "").strip()
            doi = (r.get("doi") or "").strip()
            publisher = (r.get("publisher") or "").strip()
            pages = (r.get("pages") or "").strip()
            apa = "{} ({}). {}. ".format(authors, year, title_ref)
            if source:
                if pages:
                    apa += "{}. pp. {}. ".format(source, pages)
                else:
                    apa += "{}. ".format(source)
            if publisher:
                apa += "{}. ".format(publisher)
            if doi:
                apa += "https://doi.org/{}".format(doi)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.first_line_indent = Cm(-1.25)
            p.paragraph_format.line_spacing = 2.0
            run = p.add_run(apa.strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_footer.add_run("Gerado por Cussara Academic - {}".format(datetime.now().strftime("%d/%m/%Y")))
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(128, 128, 128)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = re.sub(r'[^\w\s-]', '', title.replace(" ", "_"))[:50] + ".docx"
    return buf.getvalue(), 200, {
        "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "Content-Disposition": "attachment; filename={}".format(filename),
    }


@app.route("/api/export/<int:work_id>/pdf")
@login_required
def export_work_pdf(work_id):
    from fpdf import FPDF

    uid = session["user_id"]
    db = get_db()
    work, ctx = get_work_context(db, work_id, uid)
    if work is None:
        return jsonify({"error": "Trabalho nao encontrado."}), 404
    sections = rows_to_list(
        db.execute("SELECT * FROM sections WHERE work_id = ? ORDER BY order_index", (work_id,)).fetchall()
    )
    refs = rows_to_list(
        db.execute("SELECT * FROM references_ WHERE work_id = ? ORDER BY authors", (work_id,)).fetchall()
    )
    total_words = 0
    for s in sections:
        c = s.get("content") or ""
        s["word_count"] = len(c.split()) if c.strip() else 0
        total_words += s["word_count"]

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.set_margins(30, 25, 20)

    font_regular = os.path.join(BASE_DIR, "static", "fonts", "DejaVuSans.ttf")
    font_bold = os.path.join(BASE_DIR, "static", "fonts", "DejaVuSans-Bold.ttf")
    has_custom_font = os.path.exists(font_regular) and os.path.exists(font_bold)

    if has_custom_font:
        pdf.add_font("DejaVu", "", font_regular, uni=True)
        pdf.add_font("DejaVu", "B", font_bold, uni=True)
        regular_font = "DejaVu"
        bold_font = "DejaVu"
    else:
        regular_font = "Helvetica"
        bold_font = "Helvetica"

    pdf.add_page()
    pdf.set_font(bold_font, "B", 18)
    pdf.ln(40)
    title_text = (work.get("title") or "Trabalho Academico").upper()
    pdf.multi_cell(0, 12, title_text, align="C")
    pdf.ln(12)

    if work.get("theme"):
        pdf.set_font(regular_font, "", 12)
        pdf.cell(0, 8, "Tema: {}".format(work["theme"]), align="C", new_x="LMARGIN", new_y="NEXT")
    if work.get("area"):
        pdf.cell(0, 8, "Area: {}".format(work["area"]), align="C", new_x="LMARGIN", new_y="NEXT")
    if work.get("keywords"):
        pdf.ln(4)
        pdf.set_font(regular_font, "", 10)
        pdf.multi_cell(0, 6, "Palavras-chave: {}".format(work["keywords"]), align="C")
    pdf.ln(10)
    pdf.set_font(regular_font, "", 10)
    pdf.cell(0, 8, "Total de palavras: {}".format(total_words), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font(regular_font, "", 9)
    pdf.cell(0, 8, "Gerado por Cussara Academic - {}".format(datetime.now().strftime("%d/%m/%Y")), align="C", new_x="LMARGIN", new_y="NEXT")

    if sections and any((s.get("content") or "").strip() for s in sections):
        pdf.add_page()
        pdf.set_font(bold_font, "B", 14)
        pdf.cell(0, 10, "SUMARIO", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(6)
        pdf.set_font(regular_font, "", 11)
        for i, s in enumerate(sections):
            content = s.get("content") or ""
            if content.strip():
                pdf.cell(0, 8, "{}. {}".format(i + 1, s["title"]), new_x="LMARGIN", new_y="NEXT")

    for i, s in enumerate(sections):
        content = s.get("content") or ""
        clean = re.sub(r'<[^>]+>', ' ', content)
        clean = re.sub(r'\s+', ' ', clean).strip()
        if not clean:
            continue
        pdf.add_page()
        pdf.set_font(bold_font, "B", 14)
        pdf.cell(0, 10, "{}. {}".format(i + 1, s["title"]), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)
        pdf.set_font(regular_font, "", 12)
        paragraphs = [p.strip() for p in clean.split('\n') if p.strip()]
        for para in paragraphs:
            pdf.multi_cell(0, 7, "    " + para)
            pdf.ln(2)

    if refs:
        pdf.add_page()
        pdf.set_font(bold_font, "B", 14)
        pdf.cell(0, 10, "REFERENCIAS BIBLIOGRAFICAS", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(6)
        pdf.set_font(regular_font, "", 11)
        for r in refs:
            authors = (r.get("authors") or "").strip().rstrip(".")
            year = r.get("year") or "s.d."
            title_ref = (r.get("title") or "").strip().rstrip(".")
            source = (r.get("source") or "").strip()
            doi = (r.get("doi") or "").strip()
            publisher = (r.get("publisher") or "").strip()
            pages = (r.get("pages") or "").strip()
            apa = "{} ({}). {}. ".format(authors, year, title_ref)
            if source:
                if pages:
                    apa += "{}. pp. {}. ".format(source, pages)
                else:
                    apa += "{}. ".format(source)
            if publisher:
                apa += "{}. ".format(publisher)
            if doi:
                apa += "https://doi.org/{}".format(doi)
            x_before = pdf.get_x()
            pdf.set_x(x_before + 8)
            pdf.multi_cell(pdf.w - pdf.r_margin - pdf.get_x(), 6, apa.strip())
            pdf.ln(2)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)

    filename = re.sub(r'[^\w\s-]', '', (work.get("title") or "trabalho").replace(" ", "_"))[:50] + ".pdf"
    return buf.getvalue(), 200, {
        "Content-Type": "application/pdf",
        "Content-Disposition": "attachment; filename={}".format(filename),
    }


if __name__ == "__main__":
    with app.app_context():
        init_db()
    app.run(debug=True)
